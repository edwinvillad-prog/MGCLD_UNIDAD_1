# -*- coding: utf-8 -*-
# Unidad 4 – Super Laboratorio de Series Temporales (Streamlit)
# Autor: M.Sc. Edwin Villarreal (UPS) + ChatGPT (asistente)

from __future__ import annotations

import io
import base64
from dataclasses import dataclass
from typing import Tuple, Optional, Dict, List
from statsmodels.tsa.statespace.sarimax import SARIMAX

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import plotly.io as pio
import matplotlib.dates as mdates
import tempfile
import os

import streamlit as st
from statsmodels.tsa.seasonal import STL
from statsmodels.tsa.holtwinters import ExponentialSmoothing
from statsmodels.tsa.stattools import acf, pacf
from statsmodels.tsa.arima.model import ARIMA

st.set_page_config(
    page_title="Super Laboratorio de Series Temporales",
    page_icon="📈",
    layout="wide"   # ✅ Permite ancho completo y evita vista angosta
)

# Paleta institucional UPS
UPS_BLUE = "#002F6C"
UPS_GOLD = "#F7B500"
UPS_TEXT = "#1E1E1E"

# ===================== Configuración de página =====================

st.markdown(f"""
<h1 style='text-align:center; color:{UPS_BLUE}; font-weight:700;'>
📈 Super Laboratorio de Series Temporales
</h1>

<h4 style='text-align:center; color:black; margin-top:-10px;'>
Autor: M.Sc. Edwin Villarreal, Fís. — M.Sc. Adriana López, Fís.
</h4>

<p style='text-align:center; color:#444; font-style:italic; margin-top:-6px;'>
Maestría en Gestión de la Calidad de la Leche y sus Derivados — UPS
</p>
""", unsafe_allow_html=True)
# ===================== Utilidades =====================
def infer_frequency(index: pd.DatetimeIndex) -> Optional[str]:
    try:
        f = pd.infer_freq(index)
        if f is not None:
            return f
    except Exception:
        pass
    if len(index) >= 3:
        deltas = np.diff(index.values.astype("datetime64[ns]").astype(np.int64))
        md = np.median(deltas)
        day = 24 * 3600 * 10**9
        if np.isclose(md, day, rtol=0.05):
            return "D"
        if np.isclose(md, 7 * day, rtol=0.1):
            return "W"
        if 28 * day * 0.9 <= md <= 31 * day * 1.1:
            return "MS"
    return None


def resample_series(df: pd.DataFrame, date_col: str, value_col: str, freq: str, how: str) -> pd.DataFrame:
    out = df.set_index(date_col)[value_col].resample(freq)
    if how == "suma":
        out = out.sum()
    elif how == "promedio":
        out = out.mean()
    elif how == "último":
        out = out.last()
    else:
        out = out.mean()
    return out.to_frame(name=value_col).reset_index()


def train_test_split_series(y: pd.Series, test_size: int) -> Tuple[pd.Series, pd.Series]:
    test_size = max(1, min(test_size, max(1, len(y) // 3)))
    if len(y) <= test_size:
        test_size = max(1, len(y) // 4)
    return y.iloc[:-test_size], y.iloc[-test_size:]


def mae(y_true, y_pred):
    return float(np.mean(np.abs(np.asarray(y_true) - np.asarray(y_pred))))


def rmse(y_true, y_pred):
    return float(np.sqrt(np.mean((np.asarray(y_true) - np.asarray(y_pred)) ** 2)))


def mape(y_true, y_pred):
    y_true = np.asarray(y_true)
    y_pred = np.asarray(y_pred)
    mask = y_true != 0
    if not np.any(mask):
        return np.nan
    return float(np.mean(np.abs((y_true[mask] - y_pred[mask]) / y_true[mask])) * 100)


@dataclass
class ModelResult:
    name: str
    fit_obj: object
    yhat_in: pd.Series
    yhat_out: pd.Series
    metrics: Dict[str, float]


def fit_naive(train: pd.Series, test: pd.Series, seasonal_periods: Optional[int] = None) -> ModelResult:
    if seasonal_periods and len(train) > seasonal_periods:
        yhat_out_seed = train.iloc[-seasonal_periods:]
        reps = int(np.ceil(len(test) / seasonal_periods))
        yhat_out = pd.Series(
            np.tile(yhat_out_seed.values, reps)[: len(test)],
            index=test.index
        )
        yhat_in = train.shift(seasonal_periods)
        name = f"Naive estacional (s={seasonal_periods})"
    else:
        yhat_out = pd.Series(train.iloc[-1], index=test.index)
        yhat_in = train.shift(1)
        name = "Naive simple"

    metrics = {"MAE": mae(test, yhat_out), "RMSE": rmse(test, yhat_out), "MAPE%": mape(test, yhat_out)}
    return ModelResult(name, None, yhat_in, yhat_out, metrics)


def fit_ets(train: pd.Series, test: pd.Series, trend: Optional[str], seasonal: Optional[str], seasonal_periods: Optional[int]) -> ModelResult:
    model = ExponentialSmoothing(
        train,
        trend=trend,
        seasonal=seasonal,
        seasonal_periods=seasonal_periods if seasonal else None,
        initialization_method="estimated",
    )
    fit = model.fit(optimized=True)
    yhat_in = pd.Series(fit.fittedvalues, index=train.index)
    yhat_out = pd.Series(fit.forecast(len(test)), index=test.index)
    metrics = {"MAE": mae(test, yhat_out), "RMSE": rmse(test, yhat_out), "MAPE%": mape(test, yhat_out)}
    s_txt = f"{seasonal} s={seasonal_periods}" if seasonal else "No estacional"
    t_txt = trend if trend else "Sin tendencia"
    return ModelResult(f"ETS ({t_txt}, {s_txt})", fit, yhat_in, yhat_out, metrics)


def fit_arima(train: pd.Series, test: pd.Series, order: Tuple[int, int, int]) -> ModelResult:
    p, d, q = order
    model = ARIMA(train, order=(p, d, q))
    fit = model.fit()
    yhat_in = pd.Series(fit.fittedvalues, index=train.index)
    yhat_out = pd.Series(fit.forecast(len(test)), index=test.index)
    metrics = {"MAE": mae(test, yhat_out), "RMSE": rmse(test, yhat_out), "MAPE%": mape(test, yhat_out)}
    return ModelResult(f"ARIMA({p},{d},{q})", fit, yhat_in, yhat_out, metrics)


def fit_arima_auto(train: pd.Series, test: pd.Series,
                   max_p: int = 3, max_d: int = 2, max_q: int = 3) -> ModelResult:
    best_fit, best_order, best_aic = None, None, np.inf
    for p in range(max_p + 1):
        for d in range(max_d + 1):
            for q in range(max_q + 1):
                if p == d == q == 0:
                    continue
                try:
                    fit = ARIMA(train, order=(p, d, q)).fit()
                    if fit.aic < best_aic:
                        best_aic, best_fit, best_order = fit.aic, fit, (p, d, q)
                except Exception:
                    pass
    if best_fit is None:
        raise ValueError("ARIMA automático no convergió en el rango dado.")
    yhat_in = pd.Series(best_fit.fittedvalues, index=train.index)
    yhat_out = pd.Series(best_fit.forecast(len(test)), index=test.index)
    metrics = {"MAE": mae(test, yhat_out),
               "RMSE": rmse(test, yhat_out),
               "MAPE%": mape(test, yhat_out)}
    return ModelResult(f"ARIMA auto{best_order}", best_fit, yhat_in, yhat_out, metrics)


def fit_sarima_auto(train: pd.Series, test: pd.Series, s: int = 12,
                    max_p: int = 2, max_d: int = 1, max_q: int = 2,
                    max_P: int = 1, max_D: int = 1, max_Q: int = 1) -> ModelResult:
    best_fit, best_cfg, best_aic = None, None, np.inf
    for p in range(max_p + 1):
        for d in range(max_d + 1):
            for q in range(max_q + 1):
                for P in range(max_P + 1):
                    for D in range(max_D + 1):
                        for Q in range(max_Q + 1):
                            try:
                                model = SARIMAX(train,
                                                order=(p, d, q),
                                                seasonal_order=(P, D, Q, s),
                                                enforce_stationarity=False,
                                                enforce_invertibility=False)
                                fit = model.fit(disp=False)
                                if fit.aic < best_aic:
                                    best_aic, best_fit = fit.aic, fit
                                    best_cfg = ((p, d, q), (P, D, Q, s))
                            except Exception:
                                pass
    if best_fit is None:
        raise ValueError("SARIMA automático no convergió en el rango dado.")
    yhat_in = pd.Series(best_fit.fittedvalues, index=train.index)
    yhat_out = pd.Series(best_fit.forecast(len(test)), index=test.index)
    metrics = {"MAE": mae(test, yhat_out),
               "RMSE": rmse(test, yhat_out),
               "MAPE%": mape(test, yhat_out)}
    return ModelResult(f"SARIMA{best_cfg[0]}x{best_cfg[1]}", best_fit, yhat_in, yhat_out, metrics)


# ===================== Cartas de control =====================
def shewhart_chart(series: pd.Series) -> Tuple[pd.Series, float, float]:
    mu = float(series.mean())
    sigma = float(series.std(ddof=1))
    UCL = mu + 3 * sigma
    LCL = mu - 3 * sigma
    return pd.Series(series, name="Residuales"), UCL, LCL


def ewma_chart(series: pd.Series, lam: float = 0.2) -> Tuple[pd.Series, pd.Series, pd.Series]:
    z = series.ewm(alpha=lam, adjust=False).mean()
    sigma = float(series.std(ddof=1))
    sigma_z = sigma * np.sqrt(lam / (2 - lam))
    UCL = z * 0 + z.mean() + 3 * sigma_z
    LCL = z * 0 + z.mean() - 3 * sigma_z
    return z.rename("EWMA"), UCL.rename("UCL"), LCL.rename("LCL")


def cusum_chart(series: pd.Series, k: float, h: float) -> Tuple[pd.Series, pd.Series]:
    x = series.values
    mu0 = float(series.mean())
    cp = np.zeros(len(x))
    cm = np.zeros(len(x))
    for t in range(1, len(x)):
        cp[t] = max(0.0, x[t] - (mu0 + k) + cp[t - 1])
        cm[t] = max(0.0, (mu0 - k) - x[t] + cm[t - 1])
    return pd.Series(cp, index=series.index, name="CUSUM+"), pd.Series(cm, index=series.index, name="CUSUM-")
# ===================== Conversión de figuras =====================
def fig_to_bytes(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()
import re

import re

def resaltar_texto(texto: str) -> str:
    """
    Resalta con <b> en HTML números, porcentajes y palabras clave.
    """
    # Números
    texto = re.sub(r"(\d+(\.\d+)?)", r"<b>\1</b>", texto)
    # Porcentajes
    texto = re.sub(r"(\d+(\.\d+)?%)", r"<b>\1</b>", texto)
    # Palabras clave importantes
    palabras = ["Media", "Mediana", "Desv. Est.", "Observaciones", "tendencia", "MAPE", "RMSE", "MAE"]
    for palabra in palabras:
        texto = texto.replace(palabra, f"<b>{palabra}</b>")
    return texto


def show_conclusiones(titulo: str, lista: list):
    st.markdown(f"### 📌 {titulo}")
    for c in lista:
        c_resaltado = resaltar_texto(c)
        st.markdown(f"""
        <div style='background-color:#F9FBFD; border-left: 6px solid #4CAF50;
                    padding: 10px; margin: 8px 0; border-radius: 6px;'>
            ✅ {c_resaltado}
        </div>
        """, unsafe_allow_html=True)

def show_recomendaciones(titulo: str, lista: list):
    st.markdown(f"### 💡 {titulo}")
    for r in lista:
        r_resaltado = resaltar_texto(r)
        st.markdown(f"""
        <div style='background-color:#FFF8E1; border-left: 6px solid #FFC107;
                    padding: 10px; margin: 8px 0; border-radius: 6px;'>
            💡 {r_resaltado}
        </div>
        """, unsafe_allow_html=True)


# ===================== Barra lateral =====================
st.sidebar.header("⚙️ Parámetros")
st.sidebar.markdown("Suba un CSV con al menos una columna de **fecha** y una de **valor**.")

uploaded = st.sidebar.file_uploader("Cargar CSV", type=["csv"])
sep = st.sidebar.selectbox("Separador", [",", ";", "\t", "|"], index=0)

st.sidebar.markdown("---")
st.sidebar.subheader("Columnas y frecuencia")

# ===================== Cuerpo =====================
#st.title("📈 Unidad 4 – Super Laboratorio de Series Temporales")
st.caption("Maestría en Gestión de la Calidad de la Leche y sus Derivados – UPS")

with st.expander("📄 Formato de entrada (ejemplo)", expanded=False):
    st.write(pd.DataFrame({"fecha": ["2023-01-01", "2023-01-02", "2023-01-03"], "produccion_litros": [1200, 1180, 1215]}))
    st.markdown("Use fechas en formato ISO (YYYY-MM-DD) o con hora. Los valores deben ser numéricos.")

if uploaded is None:
    st.info("Suba un archivo para comenzar.")
    st.stop()

# ===================== Carga de datos =====================
try:
    df = pd.read_csv(uploaded, sep=sep)

    # ✅ Guardar en memoria global para otras pestañas (Markov, Control, etc.)
    st.session_state["df_base_principal"] = df.copy()
    st.session_state.setdefault("bases_cargadas", {})["base_principal.csv"] = df.copy()

except Exception as e:
    st.error(f"No se pudo leer el CSV: {e}")
    st.stop()

    
# Detectar columnas numéricas y de fecha
num_cols: List[str] = [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
date_like: List[str] = [
    c
    for c in df.columns
    if pd.api.types.is_datetime64_any_dtype(df[c]) or "date" in c.lower() or "fecha" in c.lower()
]

if not date_like:
    try:
        parsed = pd.to_datetime(df.iloc[:, 0], errors="coerce")
        if parsed.notna().mean() > 0.8:
            df.iloc[:, 0] = parsed
            date_like = [df.columns[0]]
    except Exception:
        pass

# ===================== Preparar serie temporal =====================
with st.sidebar:
    col_date = st.selectbox("Columna de fecha", options=df.columns, index=(df.columns.get_loc(date_like[0]) if date_like else 0))
    col_value = st.selectbox("Columna de valor", options=num_cols if num_cols else df.columns)
    freq_manual = st.text_input("Frecuencia (opcional, ej. D, W, MS, M, H)", value="")
    resample_how = st.selectbox("Si hay duplicados/múltiples por periodo", ["promedio", "suma", "último"])

df[col_date] = pd.to_datetime(df[col_date], errors="coerce")
df = df.dropna(subset=[col_date, col_value])
df = df.sort_values(col_date)

# Detectar si hay hora en los datos
tiene_hora = df[col_date].dt.hour.nunique() > 1

# Manejo de duplicados según configuración del usuario
st.sidebar.subheader("Manejo de fechas repetidas")
dup_method = st.sidebar.selectbox(
    "Si hay varias observaciones en la misma fecha/hora:",
    ["promedio", "suma", "último", "eliminar duplicados"]
)

if dup_method == "promedio":
    df = df.groupby(col_date, as_index=False)[col_value].mean()
elif dup_method == "suma":
    df = df.groupby(col_date, as_index=False)[col_value].sum()
elif dup_method == "último":
    df = df.groupby(col_date, as_index=False)[col_value].last()
else:
    df = df.drop_duplicates(subset=[col_date], keep="last")

if freq_manual == "D" and tiene_hora:
    st.sidebar.warning("Se detectaron horas en los datos, pero se usará frecuencia diaria (se agruparán por fecha).")
    df[col_date] = df[col_date].dt.date
    df[col_date] = pd.to_datetime(df[col_date])

if freq_manual:
    df_ts = resample_series(df, col_date, col_value, freq_manual, resample_how)
else:
    df_ts = df[[col_date, col_value]].copy()

df_ts = df_ts.set_index(col_date)
y = df_ts[col_value].asfreq(freq_manual if freq_manual else infer_frequency(df_ts.index))
y = y.dropna()

if len(y) < 10:
    st.warning("La serie tiene muy pocos datos tras limpieza. Intente con más observaciones.")
    st.stop()


# ===================== PESTAÑAS =====================
tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
    "1) Exploración",
    "2) Descomposición (STL)",
    "3) Modelado y Pronóstico",
    "4) Modelos de Markov y Monte Carlo",
    "5) Control estadístico",
    "6) Conclusiones y Recomendaciones",
    "7) Exportar Word"
])


# Acumuladores de conclusiones
conc_exploracion, conc_stl, conc_modelo, conc_control = [], [], [], []
img_series = img_forecast = img_shewhart = img_ewma = img_cusum = img_stl = None


# ===================== TAB 1: Exploración ===================== 
import matplotlib.dates as mdates
from matplotlib import animation
from matplotlib.animation import PillowWriter
import tempfile   # ✅ necesario para el GIF
import os         # ✅ necesario para manejar archivos temporales
import plotly.express as px  # ✅ para el gráfico interactivo
import plotly.io as pio
import locale

# ✅ Configuración del idioma para meses en español
try:
    locale.setlocale(locale.LC_TIME, "es_ES.utf8")  # Linux / Mac
except:
    try:
        locale.setlocale(locale.LC_TIME, "Spanish_Spain.1252")  # Windows
    except:
        pass  # Si no funciona, sigue con inglés

pio.templates.default = "plotly_white"
pio.defaults.locale = "es"   # <--- Forzar idioma español en todos los gráficos

with tab1:
    st.markdown("## 🎬 Serie temporal (exploración)")

    # Layout lado a lado: gráfica + resumen estadístico
    colA, colB = st.columns([2, 1])

    with colA:
        # Radio único con key diferente
        modo = st.radio(
            "Formato:",
            ["Estático", "GIF (ligero)"],
            horizontal=True,
            key="modo_tab1_exploracion"
        )

        # =========================================================
        # ✅ DETECCIÓN AUTOMÁTICA DE FRECUENCIA TEMPORAL
        # =========================================================
        if isinstance(y.index, pd.DatetimeIndex):
            diffs = y.index.to_series().diff().dropna()
            dias_promedio = diffs.dt.days.mean()

            if dias_promedio <= 1.5:
                freq = "Diaria"
            elif dias_promedio <= 8:
                freq = "Semanal"
            elif dias_promedio <= 32:
                freq = "Mensual"
            elif dias_promedio <= 370:
                freq = "Anual"
            else:
                freq = "Irregular"

            st.caption(f"🧭 Frecuencia detectada automáticamente: **{freq}**")
        else:
            freq = "Irregular"

        # =========================================================
        #  GRÁFICO ESTÁTICO (CON ADAPTACIÓN SEGÚN FRECUENCIA)
        # =========================================================
        if modo == "Estático":
            import plotly.graph_objects as go

            # ✅ Forzar meses en español con strftime
            fechas_es = y.index.strftime("%b %Y")

            # === Selección de estilo según frecuencia detectada ===
            # =========================================================
            # ✅ BLOQUE DE GRÁFICO SEGÚN FRECUENCIA DETECTADA
            # =========================================================
            if freq == "Diaria":
                # === Serie diaria (línea continua) ===
                fig_static = px.line(
                    x=y.index, y=y.values,
                    markers=False,
                    labels={"x": "Tiempo", "y": col_value},
                    title="Evolución diaria"
                )
                fig_static.update_traces(
                    line=dict(color="royalblue", width=2.2),
                    hovertemplate="Fecha: %{x|%d-%b-%Y}<br>Valor: %{y:.2f}"
                )
                fig_static.update_layout(
                    xaxis=dict(
                        tickformat="%b %Y",
                        tickangle=0,
                        showgrid=True,
                        ticklabelmode="period"
                    ),
                    yaxis=dict(showgrid=True, zeroline=False),
                    hovermode="x unified"
                )

            elif freq == "Semanal":
                # === Serie semanal agregada ===
                df_sem = y.resample("W").mean()
                fig_static = px.line(
                    x=df_sem.index, y=df_sem.values,
                    markers=True,
                    labels={"x": "Tiempo", "y": col_value},
                    title="Evolución semanal (serie agregada)"
                )
                fig_static.update_traces(
                    line=dict(shape="spline", color="seagreen", width=2.5),
                    marker=dict(size=5, color="seagreen", line=dict(width=0)),
                    hovertemplate="Semana: %{x|%d-%b-%Y}<br>Valor medio: %{y:.2f}"
                )
                fig_static.update_layout(
                    xaxis=dict(
                        tickformat="%d-%b %Y",
                        tickangle=0,
                        showgrid=True,
                        ticklabelmode="period"
                    ),
                    yaxis=dict(showgrid=True, zeroline=False),
                    hovermode="x unified"
                )

            elif freq == "Mensual":
                # === Serie mensual agregada ===
                df_mes = y.resample("M").mean()
                fig_static = px.line(
                    x=df_mes.index, y=df_mes.values,
                    markers=True,
                    labels={"x": "Tiempo", "y": col_value},
                    title="Evolución mensual"
                )
                fig_static.update_traces(
                    line=dict(shape="spline", color="cornflowerblue", width=2.5),
                    marker=dict(size=6, color="cornflowerblue", line=dict(width=0)),
                    hovertemplate="Mes: %{x|%b %Y}<br>Valor medio: %{y:.2f}"
                )
                fig_static.update_layout(
                    xaxis=dict(
                        tickformat="%b %Y",
                        tickangle=0,
                        showgrid=True
                    ),
                    yaxis=dict(showgrid=True, zeroline=False),
                    hovermode="x unified"
                )

            elif freq == "Anual":
                # === Serie anual agregada ===
                df_ano = y.resample("Y").mean()
                fig_static = px.line(
                    x=df_ano.index, y=df_ano.values,
                    markers=True,
                    labels={"x": "Año", "y": col_value},
                    title="Evolución anual (promedios)"
                )
                fig_static.update_traces(
                    line=dict(shape="linear", color="mediumslateblue", width=2.5),
                    marker=dict(size=7, color="mediumslateblue", line=dict(width=0)),
                    hovertemplate="Año: %{x|%Y}<br>Promedio: %{y:.2f}"
                )
                fig_static.update_layout(
                    xaxis=dict(
                        tickformat="%Y",
                        showgrid=True
                    ),
                    yaxis=dict(showgrid=True, zeroline=False),
                    hovermode="x unified"
                )

            else:
                # === Caso irregular o sin detección ===
                fig_static = px.scatter(
                    x=y.index, y=y.values,
                    labels={"x": "Tiempo", "y": col_value},
                    title="Serie irregular"
                )
                fig_static.update_traces(marker_color="firebrick", opacity=0.8)
                fig_static.update_layout(
                    xaxis=dict(showgrid=True),
                    yaxis=dict(showgrid=True),
                    hovermode="x unified"
                )

            # === Mostrar el gráfico final (ACTIVO) ===
            st.plotly_chart(fig_static, use_container_width=True)

            # 🔒 DESACTIVADO para evitar eje categórico y doble render:
            # (Se conserva el código original sin borrarlo)
            # else:
            #     fig_static = px.scatter(
            #         x=fechas_es, y=y.values,
            #         labels={"x": "Tiempo", "y": col_value},
            #         title="Serie irregular"
            #     )
            #     fig_static.update_traces(marker_color="firebrick", opacity=0.8)

            # # Hover con fecha y valor exacto (ya definido arriba por frecuencia; mantener como backup)
            # fig_static.update_traces(hovertemplate="Fecha: %{x}<br>Valor: %{y}")

            # # ✅ Layout limpio y coherente con tu diseño (versión categórica original)
            # # ⚠️ Esto fuerza eje categórico; por eso se deja desactivado.
            # # fig_static = go.Figure(fig_static)
            # # fig_static.update_layout(
            # #     xaxis=dict(
            # #         tickmode="array",
            # #         tickvals=fechas_es[::max(1, len(fechas_es)//12)],
            # #         tickangle=0
            # #     ),
            # #     hovermode="x unified"
            # # )

            # # Mostrar en Streamlit (duplicado, por eso desactivado)
            # # st.plotly_chart(fig_static, use_container_width=True)

            # === Snapshot PNG con Matplotlib (para exportar a Word) ===
            fig_tmp, ax = plt.subplots(figsize=(8, 3))
            ax.plot(y.index, y.values, marker="o", linewidth=1)
            ax.set_title("Serie temporal (snapshot estático)")
            ax.set_xlabel("Tiempo")
            ax.set_ylabel(col_value)
            ax.grid(alpha=0.3)

            ax.xaxis.set_major_locator(mdates.MonthLocator(interval=3))
            ax.xaxis.set_major_formatter(mdates.DateFormatter("%b %Y"))
            fig_tmp.autofmt_xdate(rotation=0)

            img_series = fig_to_bytes(fig_tmp)
            plt.close(fig_tmp)

        # =========================================================
        #  GRÁFICO ANIMADO (GIF)
        # =========================================================
        else:
            fig, ax = plt.subplots(figsize=(10, 3))
            ax.set_title("Serie temporal (animada)")
            ax.set_xlabel("Tiempo")
            ax.set_ylabel(col_value)
            ax.grid(alpha=0.3)

            line, = ax.plot([], [], color="steelblue", linewidth=1.5)

            ax.set_xlim(y.index.min(), y.index.max())
            ax.set_ylim(y.min() * 0.95, y.max() * 1.05)

            ax.xaxis.set_major_locator(mdates.MonthLocator(interval=4))
            ax.xaxis.set_major_formatter(mdates.DateFormatter("%b %Y"))
            fig.autofmt_xdate(rotation=0)

            def init():
                line.set_data([], [])
                return line,

            def update(frame):
                line.set_data(y.index[:frame], y.values[:frame])
                return line,

            ani = animation.FuncAnimation(
                fig, update, frames=len(y), init_func=init,
                interval=120, blit=True
            )

            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".gif")
            ani.save(tmp.name, writer=PillowWriter(fps=10))
            st.image(tmp.name, caption="Animación de la serie temporal", use_container_width=True)
            # os.remove(tmp.name)  # 🔴 No borrar aquí en Windows

    with colB:
        # === Resumen estadístico ===
        desc = y.describe(percentiles=[0.25, 0.5, 0.75])
        st.markdown("### 📑 Resumen estadístico")
        st.dataframe(desc.to_frame())
        faltantes = int(df_ts[col_value].isna().sum())

    # ================= ACF y PACF =================
    with st.expander("Autocorrelaciones (ACF/PACF)"):
        lags = min(48, max(12, len(y)//4))
        ac = acf(y, nlags=lags, fft=True)
        pc = pacf(y, nlags=min(lags, 40))

        fig_acf, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 3))
        ax1.stem(range(len(ac)), ac, basefmt=" ")
        ax1.axhline(0, color="black", linewidth=0.8)
        ax1.axhline(1.96/np.sqrt(len(y)), color="red", linestyle="--")
        ax1.axhline(-1.96/np.sqrt(len(y)), color="red", linestyle="--")
        ax1.set_title("ACF")

        ax2.stem(range(len(pc)), pc, basefmt=" ")
        ax2.axhline(0, color="black", linewidth=0.8)
        ax2.axhline(1.96/np.sqrt(len(y)), color="red", linestyle="--")
        ax2.axhline(-1.96/np.sqrt(len(y)), color="red", linestyle="--")
        ax2.set_title("PACF")

        st.pyplot(fig_acf)

    # ================= Conclusiones automáticas =================
    st.markdown("## 📌 Conclusiones (Exploración)")

    def conclusion_box(text, color, icon):
        st.markdown(
            f"""
            <div style="background-color:white; border-left: 6px solid {color};
                        padding:8px; margin:6px; border-radius:5px;
                        box-shadow: 0px 2px 4px rgba(0,0,0,0.1);">
                <b>{icon}</b> {text}
            </div>
            """, unsafe_allow_html=True
        )

    mean_val = desc["mean"]
    median_val = desc["50%"]
    min_val = desc["min"]
    max_val = desc["max"]
    q1, q3 = desc["25%"], desc["75%"]

    conclusion_box(f"La serie tiene <b>{len(y)}</b> observaciones válidas.", "#28a745", "📊")
    conclusion_box(f"La <b>media</b> es {mean_val:.2f} y la <b>mediana</b> es {median_val:.2f}.", "#0d6efd", "📈")
    conclusion_box(f"El valor <b>máximo</b> es {max_val:.2f}.", "#dc3545", "🔺")
    conclusion_box(f"El valor <b>mínimo</b> es {min_val:.2f}.", "#6f42c1", "🔻")
    conclusion_box(f"El rango intercuartílico (Q3-Q1) es {q3-q1:.2f}, con Q1={q1:.2f} y Q3={q3:.2f}.", "#17a2b8", "📐")

    if faltantes == 0:
        conclusion_box("No se detectaron valores faltantes tras la preparación de la serie.", "#28a745", "✅")
    else:
        conclusion_box(f"Se detectaron {faltantes} valores faltantes en la serie.", "#ffc107", "⚠️")

    t = np.arange(len(y))
    corr_t = np.corrcoef(t, y.values)[0, 1]
    if abs(corr_t) > 0.3:
        direction = "creciente" if corr_t > 0 else "decreciente"
        conclusion_box(f"Se observa una <b>tendencia {direction}</b> (correlación tiempo-valor |r|={corr_t:.2f}).", "#0d6efd", "📈")
    else:
        conclusion_box("No se observa una <b>tendencia global marcada</b>.", "#6c757d", "➖")


# ===================== TAB 2: Descomposición (STL) =====================
with tab2:
    import pandas as pd
    import numpy as np
    from statsmodels.tsa.seasonal import STL
    import plotly.graph_objects as go
    from plotly.subplots import make_subplots

    # ==========================================================
    # 🔍 Detección automática de frecuencia y periodo estacional
    # ==========================================================
    try:
        freq = pd.infer_freq(y.index)
    except Exception:
        freq = None

    if freq is None:
        diffs = np.median(np.diff(y.index.values).astype("timedelta64[D]").astype(int))
        if diffs <= 1:
            freq = "D"
        elif diffs <= 7:
            freq = "W"
        elif diffs <= 31:
            freq = "M"
        elif diffs <= 92:
            freq = "Q"
        else:
            freq = "A"

    if freq in ["D", "B"]:
        default_s = 7
        freq_label = "diaria"
        tick_format = "%d-%b-%Y"
        tick_interval = "M1"
    elif freq in ["W", "W-SUN", "W-MON"]:
        default_s = 52
        freq_label = "semanal"
        tick_format = "%b %Y"
        tick_interval = "M2"
    elif freq in ["M", "MS"]:
        default_s = 12
        freq_label = "mensual"
        tick_format = "%b %Y"
        tick_interval = "M3"
    elif freq in ["Q", "QS"]:
        default_s = 4
        freq_label = "trimestral"
        tick_format = "%b %Y"
        tick_interval = "M6"
    elif freq in ["A", "Y", "AS", "YS"]:
        default_s = 1
        freq_label = "anual"
        tick_format = "%Y"
        tick_interval = "Y1"
    else:
        default_s = 12
        freq_label = "no detectada"
        tick_format = "%b %Y"
        tick_interval = "M3"

    # Mensaje al usuario
    st.info(f"📆 Frecuencia detectada automáticamente: **{freq_label}** (s = {default_s})")

    seasonal_periods_default = st.number_input(
        "Periodo estacional (s)",
        min_value=2, max_value=366,
        value=default_s,
        key="stl_s"
    )

    # ==========================================================
    # ⚙️ Ajuste STL (robusto)
    # ==========================================================
    try:
        # Asegurar limpieza y orden del índice
        y = y.dropna().sort_index()

        if len(y) < int(seasonal_periods_default) * 2:
            st.warning("⚠️ La serie es demasiado corta para realizar una descomposición confiable.")
            st.stop()

        stl = STL(y, period=int(seasonal_periods_default), robust=True)
        res = stl.fit()

        # ==========================================================
        # 📊 Descomposición interactiva con Plotly
        # ==========================================================
        st.subheader("📈 Descomposición interactiva de la serie temporal")

        # --- Calcular máximos y mínimos de manera segura ---
        try:
            if freq in ["D", "B", "W"]:
                grouped = y.resample("M").mean().dropna()
            else:
                grouped = y.groupby(y.index.month).mean().dropna()

            if len(grouped) == 0:
                raise ValueError("No hay suficientes datos agrupados para identificar máximos/mínimos.")

            max_period = grouped.idxmax()
            min_period = grouped.idxmin()

            idx_max = y[y.index.isin([max_period])].idxmax()
            idx_min = y[y.index.isin([min_period])].idxmin()

        except Exception:
            # Fallback seguro: tomar máximo y mínimo global
            idx_max = y.idxmax()
            idx_min = y.idxmin()

        # Crear figura con subplots
        fig = make_subplots(
            rows=4, cols=1, shared_xaxes=True,
            subplot_titles=("Observada", "Tendencia", "Estacional", "Residuo"),
            vertical_spacing=0.06
        )

        # --- Observada ---
        fig.add_trace(
            go.Scatter(
                x=y.index, y=y.values,
                mode="lines+markers",
                name="Observada",
                line=dict(color="steelblue"),
                hovertemplate="Fecha: %{x|%d-%b-%Y}<br>Valor: %{y:.2f}"
            ), row=1, col=1
        )

        # --- Marcadores de máximo y mínimo ---
        fig.add_trace(
            go.Scatter(
                x=[idx_max], y=[y.loc[idx_max]],
                mode="markers",
                name="Máximo",
                marker=dict(color="red", size=10, symbol="circle"),
                hovertemplate="🔺 Máximo<br>Fecha: %{x|%d-%b-%Y}<br>Valor: %{y:.2f}"
            ), row=1, col=1
        )
        fig.add_trace(
            go.Scatter(
                x=[idx_min], y=[y.loc[idx_min]],
                mode="markers",
                name="Mínimo",
                marker=dict(color="blue", size=10, symbol="circle"),
                hovertemplate="🔹 Mínimo<br>Fecha: %{x|%d-%b-%Y}<br>Valor: %{y:.2f}"
            ), row=1, col=1
        )

        # --- Etiquetas visibles (sin solapamiento) ---
        fig.add_annotation(
            x=idx_max, y=y.loc[idx_max],
            text="Máximo",
            showarrow=True,
            arrowhead=2, ax=30, ay=-25,
            font=dict(color="red", size=11, family="Arial Black"),
            arrowcolor="red",
            row=1, col=1
        )
        fig.add_annotation(
            x=idx_min, y=y.loc[idx_min],
            text="Mínimo",
            showarrow=True,
            arrowhead=2, ax=30, ay=25,
            font=dict(color="blue", size=11, family="Arial Black"),
            arrowcolor="blue",
            row=1, col=1
        )

        # --- Tendencia ---
        fig.add_trace(
            go.Scatter(
                x=y.index, y=res.trend,
                mode="lines",
                name="Tendencia",
                line=dict(color="green"),
                hovertemplate="Fecha: %{x|%d-%b-%Y}<br>Tendencia: %{y:.2f}"
            ), row=2, col=1
        )

        # --- Estacional ---
        fig.add_trace(
            go.Scatter(
                x=y.index, y=res.seasonal,
                mode="lines",
                name="Estacional",
                line=dict(color="purple"),
                hovertemplate="Fecha: %{x|%d-%b-%Y}<br>Estacionalidad: %{y:.2f}"
            ), row=3, col=1
        )

        # --- Residuo ---
        fig.add_trace(
            go.Scatter(
                x=y.index, y=res.resid,
                mode="lines",
                name="Residuo",
                line=dict(color="gray"),
                hovertemplate="Fecha: %{x|%d-%b-%Y}<br>Residuo: %{y:.2f}"
            ), row=4, col=1
        )

        # --- Configuración general ---
        fig.update_layout(
            height=780,
            template="plotly_white",
            hovermode="x unified",
            showlegend=False,
            title=dict(text="🔍 Descomposición STL — Interactiva", x=0.5, xanchor="center"),
        )

        # --- Configurar eje X (menos ticks) ---
        fig.update_xaxes(
            tickformat=tick_format,
            nticks=10,
            showgrid=True,
            row=4, col=1
        )

        fig.update_yaxes(showgrid=True)

        # Mostrar figura interactiva
        st.plotly_chart(fig, use_container_width=True)

        # ==========================================================
        # 💾 Versión estática original (respaldo)
        # ==========================================================
        with st.expander("🖼️ Versión estática original (referencia)"):
            f, axes = plt.subplots(4, 1, figsize=(10, 6), sharex=True)
            axes[0].plot(y, label="Observada", color="steelblue")
            axes[0].set_title("Observada")

            axes[0].scatter(idx_max, y.loc[idx_max], color="red", s=80, zorder=5, label="Máximo")
            axes[0].scatter(idx_min, y.loc[idx_min], color="blue", s=80, zorder=5, label="Mínimo")

            axes[1].plot(res.trend, color="green"); axes[1].set_title("Tendencia")
            axes[2].plot(res.seasonal, color="purple"); axes[2].set_title("Estacional")
            axes[3].plot(res.resid, color="gray"); axes[3].set_title("Residuo")

            st.pyplot(f)

        # ------------------ Conclusiones ------------------
        st.markdown("## 📌 Conclusiones de la Descomposición STL")

        # Diccionario de meses en español (para conclusiones)
        meses_es = {
            1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril",
            5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto",
            9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
        }
        # === Variables necesarias para conclusiones (compatibilidad con versión previa) ===
        try:
            # Extraer los meses del índice máximo y mínimo
            max_month = int(y.loc[idx_max].index.month[0]) if hasattr(y.loc[idx_max], "index") else int(idx_max.month)
            min_month = int(y.loc[idx_min].index.month[0]) if hasattr(y.loc[idx_min], "index") else int(idx_min.month)

            # Valores promedio y estacionales
            max_val = y.loc[idx_max]
            min_val = y.loc[idx_min]
            estacional_max = res.seasonal.loc[idx_max]
            estacional_min = res.seasonal.loc[idx_min]
        except Exception:
            # Fallback en caso de índices irregulares
            max_month, min_month = 1, 7
            max_val, min_val = float(y.max()), float(y.min())
            estacional_max, estacional_min = 0, 0


        def conclusion_box(text: str, color: str = "#28a745", emoji: str = "✅"):
            st.markdown(
                f"""
                <div style="background-color:#ffffff; border-left:6px solid {color};
                            box-shadow: 0px 2px 6px rgba(0,0,0,0.1);
                            padding:10px; margin:10px 0; border-radius:6px; font-size:15px;">
                    <b style="font-size:16px">{emoji}</b> {text}
                </div>
                """,
                unsafe_allow_html=True,
            )

        # Interpretación automática
        conclusion_box("La serie <b>observada</b> refleja el comportamiento original (tendencia, estacionalidad y ruido).", "#28a745", "📊")

        if res.trend.dropna().iloc[-1] > res.trend.dropna().iloc[0]:
            conclusion_box("La <b>tendencia</b> es <b>creciente</b>: el promedio de la serie va subiendo con el tiempo.", "#0d6efd", "📈")
        else:
            conclusion_box("La <b>tendencia</b> es <b>decreciente</b>: el promedio de la serie va bajando con el tiempo.", "#0d6efd", "📉")

        if np.var(res.seasonal) > 0.05 * np.var(y.values):
            conclusion_box("Existe una <b>estacionalidad fuerte</b>, con subidas y bajadas que se repiten regularmente.", "#17a2b8", "🔄")
        else:
            conclusion_box("La <b>estacionalidad</b> es débil o moderada en el periodo analizado.", "#17a2b8", "🔄")

        if np.std(res.resid) < 0.2 * np.std(y.values):
            conclusion_box("Los <b>residuos</b> son pequeños → el modelo STL capturó bien la estructura.", "#6f42c1", "🟣")
        else:
            conclusion_box("Los <b>residuos</b> muestran alta variabilidad → STL no explica toda la dinámica.", "#6f42c1", "🟣")

        conclusion_box("Los picos en el <b>residuo</b> pueden indicar <b>anomalías o choques puntuales</b> en la serie.", "#dc3545", "⚠️")

        # Máximo y mínimo con valores y estacionalidad
        conclusion_box(
            f"🔴 El mes con <b>mayor valor promedio</b> es <b>{meses_es[max_month]}</b> "
            f"con <b>{max_val:.2f}</b>, impulsado por una estacionalidad de <b style='color:red;'>+{estacional_max:.2f}</b>.",
            "#dc3545", "📅"
        )

        conclusion_box(
            f"🔵 El mes con <b>menor valor promedio</b> es <b>{meses_es[min_month]}</b> "
            f"con <b>{min_val:.2f}</b>, afectado por una estacionalidad de <b style='color:blue;'>{estacional_min:.2f}</b>.",
            "#0d6efd", "📅"
        )

        # ------------------ Resumen numérico ------------------
        var_total = np.var(y.values, ddof=1)
        var_trend = np.var(res.trend.dropna().values, ddof=1)
        var_seas  = np.var(res.seasonal.dropna().values, ddof=1)
        var_res   = np.var(res.resid.dropna().values, ddof=1)

        def ratio(v): 
            return (100*v/var_total) if var_total > 0 else 0.0

        st.markdown("### 📊 Resumen numérico de la descomposición")
        st.markdown(
            f"""
            <div style="background-color:#f1f3f5; border-left:6px solid #6c757d;
                        padding:15px; margin:12px 0; border-radius:8px; font-size:15px;">
                <b>Composición de varianza:</b><br>
                • Tendencia: <b style="color:#0d6efd;">{ratio(var_trend):.1f}%</b><br>
                • Estacional: <b style="color:#17a2b8;">{ratio(var_seas):.1f}%</b><br>
                • Residuo: <b style="color:#dc3545;">{ratio(var_res):.1f}%</b>
            </div>
            """,
            unsafe_allow_html=True,
        )
        # -------------------------------------------------------------
        # 💾 Guardar versión estática de la descomposición STL en sesión
        # -------------------------------------------------------------
        import io
        buf = io.BytesIO()
        f.savefig(buf, format="png", bbox_inches="tight")
        st.session_state["img_stl"] = buf.getvalue()

    except Exception as e:
        st.warning(f"STL no pudo ajustarse: {e}")
        st.warning("No fue posible ajustar STL con los parámetros actuales.")


# ===================== TAB 3: Modelado y Pronóstico ===================== 
with tab3:
    # ===================== Sección teórica general =====================
    with st.expander("📘 Fundamento teórico de los modelos de series temporales", expanded=False):

        # ====================================================
        # 1️⃣ MODELO ARIMA
        # ====================================================
        st.markdown("## 1️⃣ Modelo ARIMA (AutoRegressive Integrated Moving Average)")

        st.markdown("""
        El modelo **ARIMA(p, d, q)** combina tres componentes fundamentales:

        - **AR (p):** parte *autoregresiva*, modela la dependencia lineal con los valores pasados de la serie.  
        - **I (d):** parte *integrada*, elimina la tendencia y vuelve la serie estacionaria.  
        - **MA (q):** parte de *promedio móvil*, modela el ruido o error pasado.
        """)

        st.markdown("**Ecuación general:**")
        st.latex(r"""
        y_t = c + \sum_{i=1}^{p} \phi_i\,y_{t-i} + \sum_{j=1}^{q} \theta_j\,\varepsilon_{t-j} + \varepsilon_t
        """)

        st.markdown("**Donde:**")
        st.markdown("""
        - $y_t$: valor observado en el tiempo $t$.  
        - $\\varepsilon_t$: error aleatorio (*ruido blanco*).  
        - $\\phi_i$, $\\theta_j$: coeficientes autoregresivos y de medias móviles.  
        """)

        st.info("""
        💡 **Aplicación:** En el contexto de la **calidad de la leche**, el ARIMA permite pronosticar variables como el **conteo microbiano**, la **temperatura del tanque** o el **precio del productor**, considerando **tendencias y fluctuaciones a corto plazo**.
        """)

        st.divider()

        # ====================================================
        # 2️⃣ MODELO SARIMA
        # ====================================================
        st.markdown("## 2️⃣ Modelo SARIMA (Seasonal ARIMA)")

        st.markdown("""
        El modelo **SARIMA(p, d, q)(P, D, Q, s)** extiende a ARIMA incorporando **estacionalidad**, con un componente adicional que captura patrones repetitivos cada $s$ períodos.
        """)

        st.markdown("**Ecuación general:**")
        st.latex(r"""
        \Phi_P(L^s)\,\phi_p(L)\,(1 - L)^d\,(1 - L^s)^D\,y_t =
        \Theta_Q(L^s)\,\theta_q(L)\,\varepsilon_t
        """)

        st.markdown("**Donde:**")
        st.markdown("""
        - $s$: número de observaciones por ciclo (por ejemplo, 12 meses o 52 semanas).  
        - $\\Phi_P$, $\\Theta_Q$: componentes estacionales autoregresivos y de medias móviles.  
        """)

        st.info("""
        💡 **Aplicación:** Útil para analizar variaciones **anuales o semanales** en el **conteo de UFC/ml**, reflejando patrones de contaminación según la estación seca o lluviosa.
        """)

        st.divider()

        # ====================================================
        # 3️⃣ MODELO ETS
        # ====================================================
        st.markdown("## 3️⃣ Modelo ETS (Error–Trend–Seasonal — Holt–Winters)")

        st.markdown("""
        Los modelos **ETS** utilizan *suavizamiento exponencial* para capturar patrones de **nivel**, **tendencia** y **estacionalidad**.  
        Se ajustan mediante parámetros $\\alpha$, $\\beta$, $\\gamma$ que determinan el peso de las observaciones recientes.
        """)

        st.markdown("**Ecuaciones del modelo aditivo:**")
        st.latex(r"""
        \begin{cases}
        \ell_t = \alpha (y_t - s_{t-m}) + (1 - \alpha)(\ell_{t-1} + b_{t-1}) \\
        b_t = \beta (\ell_t - \ell_{t-1}) + (1 - \beta)b_{t-1} \\
        s_t = \gamma (y_t - \ell_t) + (1 - \gamma)s_{t-m}
        \end{cases}
        """)

        st.markdown("**Donde:**")
        st.markdown("""
        - $\\ell_t$: nivel suavizado.  
        - $b_t$: tendencia estimada.  
        - $s_t$: componente estacional.  
        - $m$: longitud del ciclo.  
        """)

        st.info("""
        💡 **Aplicación:** En la **producción de leche** o el **consumo energético mensual**, el modelo ETS representa con precisión **tendencias graduales y estacionalidad leve**.
        """)

        st.divider()

        # ====================================================
        # 4️⃣ MODELO NAÏVE
        # ====================================================
        st.markdown("## 4️⃣ Modelo Naïve (Referencia estacional)")

        st.markdown("""
        El modelo **Naïve** asume que el futuro repetirá el pasado sin cambios estructurales.  
        Es la base de comparación para validar modelos más complejos.
        """)

        st.markdown("**Ecuaciones:**")
        st.latex(r"""
        \hat{y}_{t+h} = y_t \quad \text{o bien} \quad \hat{y}_{t+h} = y_{t+h-s}
        """)

        st.markdown("**Donde:**")
        st.markdown("""
        - $\\hat{y}_{t+h}$: valor pronosticado $h$ pasos adelante.  
        - $y_t$: valor observado actual.  
        - $s$: periodo estacional de repetición.  
        """)

        st.info("""
        💡 **Aplicación:** Útil para estimaciones rápidas de **producción semanal** o **eficiencia energética**, donde el proceso se mantiene estable o sin tendencia significativa.
        """)

        st.divider()

        # ====================================================
        # 🧠 COMPARATIVO FINAL
        # ====================================================
        st.markdown("### 🧠 Comparativo general de los modelos")

        st.markdown("""
        | Modelo | Tendencia | Estacionalidad | Adecuado para |
        |:--------|:----------:|:--------------:|:--------------|
        | **ARIMA** | ✅ | ❌ | Series con tendencia sin estacionalidad |
        | **SARIMA** | ✅ | ✅ | Series con patrones repetitivos (mensuales o anuales) |
        | **ETS (Holt–Winters)** | ✅ | ✅ | Series con tendencia y estacionalidad suave |
        | **Naïve** | ⚠️ Simple | ✅ | Procesos estables o como línea base |

        💬 *En la gestión de la calidad de la leche, estos modelos permiten anticipar desviaciones en parámetros críticos, optimizar mantenimiento y reducir pérdidas por contaminación.*
        """, unsafe_allow_html=True)



    # ===================== Detección automática de frecuencia =====================
    import pandas as pd
    import numpy as np
    from statsmodels.tsa.stattools import acf

    # Intentar detectar frecuencia automáticamente
    try:
        freq = pd.infer_freq(y.index)
    except Exception:
        freq = None

    # =====================================================
    # 🔍 Mejora: detección automática de frecuencia y rangos
    # =====================================================
    # Si no se detecta frecuencia explícita, intentamos inferirla
    # según la mediana de la diferencia de fechas
    if freq is None:
        diffs = np.median(np.diff(y.index.values).astype('timedelta64[D]').astype(int))
        if diffs <= 1:
            freq = "D"
        elif diffs <= 7:
            freq = "W"
        elif diffs <= 31:
            freq = "M"
        elif diffs <= 92:
            freq = "Q"
        else:
            freq = "A"

    # =====================================================
    # 🎯 Ajuste dinámico de parámetros según la frecuencia
    # =====================================================
    if freq in ["D", "B"]:   # Diario o días hábiles
        s_default = 7    # semanal
        p_range, d_range, q_range = (8, 2, 8)
    elif freq in ["W", "W-SUN", "W-MON"]:  # Semanal
        s_default = 52   # anual en semanas
        p_range, d_range, q_range = (4, 2, 4)
    elif freq in ["M", "MS"]:  # Mensual
        s_default = 12   # anual
        p_range, d_range, q_range = (4, 2, 4)
    elif freq in ["Q", "QS"]:  # Trimestral
        s_default = 4
        p_range, d_range, q_range = (2, 1, 2)
    elif freq in ["A", "Y", "AS", "YS"]:  # Anual
        s_default = 1
        p_range, d_range, q_range = (1, 1, 1)
    else:
        # Fallback si no se detecta
        s_default = 12
        p_range, d_range, q_range = (4, 2, 4)

    # =====================================================
    # 🧠 Análisis adicional: sugerencia de p y q mediante ACF
    # =====================================================
    try:
        ac = acf(y, nlags=min(24, len(y)//3), fft=True)
        lag1, lag2 = np.argsort(-np.abs(ac[1:5]))[:2] + 1
        p_sugerido = int(lag1)
        q_sugerido = int(lag2)
        from statsmodels.tsa.stattools import adfuller

        # --- Test de estacionariedad para sugerir d ---
        try:
            adf_pvalue = adfuller(y.dropna())[1]
            if adf_pvalue < 0.05:
                d_sugerido = 0  # ya estacionaria
            else:
                d_sugerido = 1  # necesita al menos una diferenciación
        except Exception:
            d_sugerido = 1

    except Exception:
        p_sugerido, q_sugerido = 1, 1

    # Guardar en session_state para que otras tabs lo usen
    st.session_state["freq_detected"] = freq
    st.session_state["s_default"] = s_default
    st.session_state["p_range"] = p_range
    st.session_state["d_range"] = d_range
    st.session_state["q_range"] = q_range

    st.info(f"📅 Frecuencia detectada: **{freq}** — Período estacional estimado: s={s_default}")

    # ===================== Inferencia automática de frecuencia =====================
    try:
        freq = pd.infer_freq(y.index)
    except Exception:
        freq = None

    if freq in ["M", "MS"]:      # Mensual
        seasonal_periods_default = 12
    elif freq in ["D"]:          # Diario
        seasonal_periods_default = 7
    elif freq in ["Q", "QS"]:    # Trimestral
        seasonal_periods_default = 4
    elif freq in ["A", "AS"]:    # Anual
        seasonal_periods_default = 1
    else:
        seasonal_periods_default = 12  # fallback mensual

    col1, col2, col3 = st.columns(3)
    with col1:
        horizon = st.number_input(
            "Horizonte de pronóstico (pasos)", 
            min_value=1, max_value=1000, 
            value=min(12, max(1, len(y)//4)), 
            key="hor"
        )
    with col2:
        test_size = st.number_input(
            "Tamaño del conjunto de prueba", 
            min_value=1, max_value=max(1, len(y)//2), 
            value=max(1, int(0.2 * len(y))), 
            key="ts"
        )
    with col3:
        seasonal_periods = st.number_input(
            "s (para ETS)", 
            min_value=2, max_value=366, 
            value=int(seasonal_periods_default), 
            key="s_ets"
        )

    train, test = train_test_split_series(y, int(test_size))
    models: List[ModelResult] = []

    models.append(fit_naive(train, test, seasonal_periods=int(seasonal_periods)))

    try:
        models.append(fit_ets(train, test, trend=None, seasonal=None, seasonal_periods=None))
    except Exception as e:
        st.warning(f"SES falló: {e}")

    for tr in [None, "add"]:
        for seas in [None, "add", "mul"]:
            if seas is None and tr is None:
                continue
            try:
                models.append(fit_ets(train, test, trend=tr, seasonal=seas, seasonal_periods=int(seasonal_periods)))
            except Exception:
                pass

    # =====================================================
    # 🔹 Bloque ARIMA mejorado con detección automática
    # =====================================================
    with st.expander("ARIMA (opcional)"):
        p = st.slider("p", 0, st.session_state["p_range"],
                    min(p_sugerido, st.session_state["p_range"]), key="p")
        d = st.slider("d", 0, st.session_state["d_range"],
                    min(d_sugerido, st.session_state["d_range"]), key="d")
        q = st.slider("q", 0, st.session_state["q_range"],
                    min(q_sugerido, st.session_state["q_range"]), key="q")

        st.caption(
            f"🧠 Rango ARIMA detectado: p∈[0,{p_range}], d∈[0,{d_range}], q∈[0,{q_range}] (s={s_default})"
        )
        st.caption(
            f"📊 Sugerencia automática: p≈{p_sugerido}, d≈{d_sugerido}, q≈{q_sugerido} (según ACF + ADF)"
        )


        try_arima = st.checkbox("Añadir ARIMA a la comparación", value=False, key="do_arima")
        if try_arima:
            try:
                models.append(fit_arima(train, test, (p, d, q)))
            except Exception as e:
                st.warning(f"ARIMA falló: {e}")

    # ===================== Búsqueda automática de modelos =====================
    st.markdown("---")
    st.markdown("### 🔎 Búsqueda automática de modelos (ARIMA / SARIMA)")

    colA, colB = st.columns(2)
    with colA:
        use_arima_auto = st.checkbox("Probar ARIMA automático", value=False, key="do_arima_auto")
        max_p_auto = st.slider("max p (ARIMA auto)", 0, st.session_state["p_range"], 3, key="p_auto")
        max_d_auto = st.slider("max d (ARIMA auto)", 0, st.session_state["d_range"], 1, key="d_auto")
        max_q_auto = st.slider("max q (ARIMA auto)", 0, st.session_state["q_range"], 3, key="q_auto")

    with colB:
        use_sarima_auto = st.checkbox("Probar SARIMA automático", value=False, key="do_sarima_auto")
        s_season = st.number_input("Periodo estacional (s) para SARIMA", min_value=2, max_value=366,
                                   value=int(seasonal_periods), key="s_sarima")

    if use_arima_auto:
        try:
            models.append(
                fit_arima_auto(train, test, max_p=max_p_auto, max_d=max_d_auto, max_q=max_q_auto)
            )
        except Exception as e:
            st.warning(f"ARIMA automático falló: {e}")

    if use_sarima_auto:
        try:
            models.append(fit_sarima_auto(train, test, s=int(s_season)))
        except Exception as e:
            st.warning(f"SARIMA automático falló: {e}")

    # ===================== Comparación de modelos =====================
    if not models:
        st.warning("No se logró ajustar ningún modelo.")
        st.stop()

    st.subheader("📊 Comparación de modelos (menor es mejor)")

    tbl = (
        pd.DataFrame([{**m.metrics, "Modelo": m.name} for m in models])
        .set_index("Modelo")
        .sort_values("RMSE")
    )

    def map_equivalente(nombre):
        nombre = nombre.lower()
        if "add, mul" in nombre:
            return "Holt–Winters multiplicativo"
        elif "add, add" in nombre:
            return "Holt–Winters aditivo"
        elif "sin tendencia" in nombre and "mul" in nombre:
            return "SES estacional multiplicativo"
        elif "sin tendencia" in nombre and "add" in nombre:
            return "SES estacional aditivo"
        elif "add" in nombre and "no estacional" in nombre:
            return "Holt (doble suavizamiento)"
        elif "sin tendencia" in nombre and "no estacional" in nombre:
            return "SES (simple)"
        elif "naive" in nombre:
            return "Modelo Naïve estacional"
        else:
            return "—"

    tbl["Equivalente clásico"] = tbl.index.to_series().apply(map_equivalente)
    st.dataframe(
        tbl[["Equivalente clásico", "MAE", "RMSE", "MAPE%"]]
            .style.format({"MAE": "{:.3f}", "RMSE": "{:.3f}", "MAPE%": "{:.2f}"})
    )

    best = tbl.index[0]
    st.success(f"🏆 Mejor en prueba: **{best}**")
    best_model = next(m for m in models if m.name == best)

    # Guardar para otras pestañas
    st.session_state["tbl"] = tbl
    st.session_state["best"] = best


    # ===================== Gráfico ajuste vs. prueba =====================
    
    st.subheader("🎯 Desempeño del modelo")
    f3, ax3 = plt.subplots(figsize=(11, 3))
    ax3.plot(train.index, train.values, label="Train")
    ax3.plot(train.index, best_model.yhat_in.reindex(train.index), label="Ajuste")
    ax3.plot(test.index, test.values, label="Test")
    ax3.plot(test.index, best_model.yhat_out.reindex(test.index), label="Pronóstico (test)")
    ax3.legend()
    ax3.set_title(f"Desempeño del modelo {best}: Ajuste (Train) y Pronóstico (Test)")
    st.pyplot(f3)
    img_fit = fig_to_bytes(f3)
    # ===================== Pronóstico futuro con el mejor modelo ===================== 
    st.subheader("🔮 Pronóstico futuro")

    # Reentrenar según el mejor modelo encontrado
    final_fit = None
    future = None

    try:
        if best.startswith("ARIMA auto") or best.startswith("ARIMA("):
            # ARIMA manual o automático
            if "auto" in best:
                order = best_model.fit_obj.model_orders.get('ar'), best_model.fit_obj.model_orders.get('diff'), best_model.fit_obj.model_orders.get('ma')
                try:
                    ord_text = str(best).replace("ARIMA auto", "").strip()
                    order = eval(ord_text)
                except Exception:
                    pass
                final_fit = ARIMA(y, order=order).fit()
            else:
                final_fit = ARIMA(y, order=(p, d, q)).fit()
            future = final_fit.forecast(int(horizon))

        elif best.startswith("SARIMA"):
            # Parseo de órdenes desde el nombre
            try:
                tag = best.replace("SARIMA", "")
                non_seas, seas = tag.split("x")
                order = eval(non_seas)
                seas_order = eval(seas)
            except Exception:
                order = (1, 1, 1)
                seas_order = (1, 1, 1, int(s_season))
            final_fit = SARIMAX(y, order=order, seasonal_order=seas_order,
                            enforce_stationarity=False, enforce_invertibility=False).fit(disp=False)
            future = final_fit.forecast(int(horizon))

        elif best.startswith("ETS"):
            trend = "add" if "add" in best else None
            seasonal = "mul" if "mul" in best else ("add" if "add" in best else None)
            final_fit = ExponentialSmoothing(
                y, trend=trend, seasonal=seasonal, seasonal_periods=int(seasonal_periods),
                initialization_method="estimated"
            ).fit()
            future = final_fit.forecast(int(horizon))

        else:  # Naive
            if "estacional" in best:
                last_season = y.iloc[-int(seasonal_periods):]
                reps = int(np.ceil(int(horizon) / int(seasonal_periods)))
                future_vals = np.tile(last_season.values, reps)[: int(horizon)]
            else:
                future_vals = np.repeat(y.iloc[-1], int(horizon))
            future_index = pd.date_range(y.index[-1], periods=int(horizon) + 1, freq=y.index.freq)[1:]
            future = pd.Series(future_vals, index=future_index)
            final_fit = None
    except Exception as e:
        st.error(f"No se pudo generar el pronóstico futuro con el modelo seleccionado: {e}")

# Gráfico de pronóstico
# ===================== Pronóstico futuro con el mejor modelo =====================

    if future is not None and len(future) > 0:
        import plotly.graph_objects as go

        # Crear figura interactiva
        fig_forecast = go.Figure()

        # ======================================================
        # 🧭 Detección automática de frecuencia y formato del eje
        # ======================================================
        try:
            freq = pd.infer_freq(y.index)
        except Exception:
            freq = None

        # Definir formato de fechas y espaciado de ticks según la frecuencia
        if freq in ["D", "B"]:
            formato_fecha = "%d-%b-%Y"
            titulo_eje_x = "Tiempo (diario)"
            dtick_val = 7  # cada semana
        elif freq in ["W", "W-SUN", "W-MON"]:
            formato_fecha = "%d-%b-%Y"
            titulo_eje_x = "Tiempo (semanal)"
            dtick_val = 4  # 4 semanas ≈ 1 mes
        elif freq in ["M", "MS"]:
            formato_fecha = "%b %Y"
            titulo_eje_x = "Tiempo (mensual)"
            dtick_val = 1
        elif freq in ["Q", "QS"]:
            formato_fecha = "T%q-%Y"
            titulo_eje_x = "Tiempo (trimestral)"
            dtick_val = 1
        elif freq in ["A", "Y", "AS", "YS"]:
            formato_fecha = "%Y"
            titulo_eje_x = "Tiempo (anual)"
            dtick_val = 1
        else:
            formato_fecha = "%b %Y"
            titulo_eje_x = "Tiempo"
            dtick_val = 1

        # Fechas formateadas para histórico y futuro
        fechas_hist = y.index.strftime(formato_fecha)
        fechas_fut = future.index.strftime(formato_fecha)

        # Serie observada
        fig_forecast.add_trace(go.Scatter(
            x=fechas_hist, y=y.values,
            mode="lines+markers",
            name="Observado",
            line=dict(color="steelblue"),
            hovertemplate="Fecha: %{x}<br>Valor: %{y:.2f}"
        ))

        # Pronóstico futuro
        fig_forecast.add_trace(go.Scatter(
            x=fechas_fut, y=future.values,
            mode="lines+markers",
            name="Pronóstico",
            line=dict(color="darkorange", dash="dash"),
            hovertemplate="Fecha: %{x}<br>Proyección: %{y:.2f}"
        ))

        # ======================================================
        # 🎨 Ajuste visual del eje x según frecuencia y tamaño
        # ======================================================
        # Reducir el número de etiquetas de fecha para series largas
        num_points = len(y)
        max_ticks = 12  # máximo de etiquetas legibles

        # Ajuste automático del paso de ticks según el tamaño de la serie
        if num_points <= max_ticks:
                tickstep = 1
        else:
            tickstep = max(1, num_points // max_ticks)

        fig_forecast.update_layout(
            title=f"Pronóstico con {best}",
            xaxis=dict(
                title=titulo_eje_x,
                tickangle=-30,          # inclinación ligera para mejor lectura
                tickmode="array",       # modo de ticks definidos manualmente
                tickvals=list(range(0, len(fechas_hist), tickstep)),
                ticktext=[fechas_hist[i] for i in range(0, len(fechas_hist), tickstep)],
                showgrid=True
            ),
            yaxis=dict(title=col_value, showgrid=True),
            hovermode="x unified"
        )


        # Mostrar en la app
        st.plotly_chart(fig_forecast, use_container_width=True)

        # Guardar snapshot PNG para exportar a Word
        f4, ax4 = plt.subplots(figsize=(11, 3))
        ax4.plot(y.index, y.values, label="Observado")
        ax4.plot(future.index, future.values, label="Pronóstico", linestyle="--", color="orange")
        ax4.legend()
        ax4.set_title(f"Pronóstico {best}")
        img_forecast = fig_to_bytes(f4)
        plt.close(f4)

    else:
        img_forecast = None
        st.warning("No se pudo graficar el pronóstico futuro.")



# ===================== Conclusiones del modelado =====================
    st.markdown("## 📌 Conclusiones del Modelado y Pronóstico")

    def conclusion_box(text, color, icon):
        st.markdown(
            f"""
            <div style="background-color:white; border-left: 6px solid {color};
                        padding:8px; margin:6px; border-radius:5px;
                        box-shadow: 0px 2px 4px rgba(0,0,0,0.1);">
                <b>{icon}</b> {text}
            </div>
            """, unsafe_allow_html=True
        )

# ===================== 1. Mejor modelo y métricas =====================
    m = tbl.loc[best]
    conclusion_box(
        f"🏆 El mejor modelo fue <b>{best}</b> con "
        f"<b>RMSE={m['RMSE']:.3f}</b>, <b>MAE={m['MAE']:.3f}</b>, "
        f"<b>MAPE={m['MAPE%']:.2f}%</b>.",
         "#28a745", ""
    )

# ===================== 2. Precisión =====================
    if pd.notna(m["MAPE%"]) and m["MAPE%"] < 10:
        conclusion_box("Alta precisión (MAPE < 10%). Puede usarse para decisiones operativas con confianza.", "#0d6efd", "✅")
    elif pd.notna(m["MAPE%"]) and m["MAPE%"] < 20:
        conclusion_box("Precisión moderada (10–20%). Recomendable validación con criterio experto.", "#ffc107", "⚠️")
    else:
        conclusion_box("Precisión baja (MAPE ≥ 20%). Se sugiere mejorar calidad de datos o probar modelos adicionales.", "#dc3545", "❌")

# ===================== 3. Robustez =====================
    if len(test) < 5:
        conclusion_box("⚠️ El conjunto de prueba es pequeño → las métricas podrían no ser representativas.", "#ffc107", "ℹ️")

# ===================== 4. Tipo de modelo =====================
    if best.startswith("ETS") and "add" in best:
        conclusion_box("El modelo ETS detectó <b>tendencia aditiva</b> en la serie.", "#6f42c1", "📐")
    elif best.startswith("SARIMA"):
        conclusion_box("El modelo SARIMA aprovechó la <b>estacionalidad</b> detectada en los datos.", "#6610f2", "🔎")
    elif "naive" in best.lower():
        conclusion_box("El modelo Naive fue el mejor; esto puede indicar <b>estacionalidad fuerte</b> o <b>poca variabilidad</b>.", "#dc3545", "ℹ️")

# ===================== 5. Comparación con el histórico =====================
    if st.session_state.get("future") is not None:
        fmean = float(st.session_state["future"].mean())
        ymean = float(y.mean())
        if fmean > ymean:
            conclusion_box("El pronóstico promedio está por <b>encima</b> del histórico → <b>tendencia al alza</b>.", "#17a2b8", "📈")
        elif fmean < ymean:
            conclusion_box("El pronóstico promedio está por <b>debajo</b> del histórico → <b>tendencia a la baja</b>.", "#17a2b8", "📉")
        else:
            conclusion_box("El pronóstico promedio es <b>similar</b> al histórico → estabilidad.", "#6c757d", "➖")

# ===================== 6. Recomendaciones prácticas =====================
    if "ETS" in best:
        conclusion_box("🛠️ El modelo ETS es adecuado si se espera <b>continuidad en estacionalidad</b> y <b>tendencia suave</b>.", "#6f42c1", "")
    elif "ARIMA" in best or "SARIMA" in best:
        conclusion_box("🛠️ El modelo ARIMA/SARIMA es más robusto para <b>estructuras temporales complejas</b> (ej. shocks, ciclos largos).", "#6610f2", "")
    elif "Naive" in best:
        conclusion_box("🛠️ El modelo Naive puede usarse como <b>línea base</b>, pero se recomienda probar alternativas más sofisticadas para decisiones críticas.", "#dc3545", "")

    if future is not None and len(future) > 0:
        fmean = float(future.mean())
        ymean = float(y.mean())

        # Texto con números resaltados
        text_mean = f"📊 El valor promedio proyectado es <b>{fmean:.2f}</b>, comparado con el histórico de <b>{ymean:.2f}</b>."
        conc_modelo.append(f"El valor promedio proyectado es {fmean:.2f}, comparado con el histórico de {ymean:.2f}.")
        conclusion_box(text_mean, "#20c997", "")

        # Tendencia resaltada
        if fmean > ymean * 1.05:
            text_growth = "📈 El modelo proyecta un <b style='color:green;'>crecimiento sostenido</b> en comparación al histórico."
            conc_modelo.append("El modelo proyecta un crecimiento sostenido en comparación al histórico.")
            conclusion_box(text_growth, "#0d6efd", "")

        elif fmean < ymean * 0.95:
            text_drop = "📉 El modelo proyecta una <b style='color:red;'>disminución significativa</b> en comparación al histórico."
            conc_modelo.append("El modelo proyecta una disminución significativa en comparación al histórico.")
            conclusion_box(text_drop, "#dc3545", "")

        else:
            text_stable = "➖ El modelo proyecta <b>estabilidad</b> en los próximos períodos."
            conc_modelo.append("El modelo proyecta estabilidad en los próximos períodos.")
            conclusion_box(text_stable, "#6c757d", "")

        # Estacionalidad resaltada
        if "ETS" in best or "SARIMA" in best:
            text_season = "🔄 Se espera que la <b>estacionalidad</b> observada se mantenga en el futuro."
            conc_modelo.append("Se espera que la estacionalidad observada se mantenga en el futuro.")
            conclusion_box(text_season, "#17a2b8", "")

        # Guardar el modelo final y los pronósticos para uso posterior
        st.session_state["final_fit"] = final_fit
        st.session_state["future"] = future
        st.session_state["best_model_name"] = best
        st.session_state["seasonal_periods_used"] = int(seasonal_periods)
    else:
        st.warning("⚠️ No se generó pronóstico futuro; no se guardaron resultados para control estadístico.")


# ===========================================================
# 🔁 TAB 4 — MODELOS DE MARKOV + SIMULACIÓN MONTE CARLO
# ===========================================================
with tab4:
    st.header("🔁 Modelos de Markov y Simulación Monte Carlo")
    st.markdown("Esta sección estima la **matriz de transición de un proceso de Markov** a partir de los datos históricos "
                "y simula su evolución mediante el método de **Monte Carlo**, "
                "para analizar la probabilidad de permanecer o cambiar entre estados.")

    # ======================================================
    # 💾 Límite de carga
    # (Configúralo fuera del script: CLI o config.toml)
    # ======================================================
    # st.set_option("server.maxUploadSize", 1000)  # ← PROVOCA EXCEPCIÓN EN TU VERSIÓN

    # ======================================================
    # 🧠 Detección automática de base compatible (todas las cargadas)
    # ======================================================
    df_auto = None
    nombre_base = None

    # Buscar entre todas las bases cargadas
    if "bases_cargadas" in st.session_state:
        for nombre, df_temp in st.session_state["bases_cargadas"].items():
            df_temp.columns = [c.strip() for c in df_temp.columns]
            columnas_lower = [c.lower() for c in df_temp.columns]
            columnas_estado = [df_temp.columns[i] for i, c in enumerate(columnas_lower) if "estado" in c]

            if len(columnas_estado) >= 2:
                col_estado_actual = columnas_estado[0]
                col_estado_siguiente = columnas_estado[1]
                st.success(f"✅ Detectadas columnas de transición en **{nombre}**: {col_estado_actual}, {col_estado_siguiente}")
                df_auto = df_temp.rename(columns={
                    col_estado_actual: "Estado_Actual",
                    col_estado_siguiente: "Estado_Siguiente"
                })
                nombre_base = nombre
                break

            elif len(columnas_estado) == 1:
                col_estado = columnas_estado[0]
                st.info(f"ℹ️ Detectada columna de estado en **{nombre}**: {col_estado} — se usará para calcular transiciones consecutivas.")
                df_auto = df_temp.rename(columns={col_estado: "Estado"})
                nombre_base = nombre
                break

    if df_auto is not None:
        st.markdown(f"<p style='color:#0d6efd;font-weight:bold;'>📂 Base seleccionada automáticamente: {nombre_base}</p>", unsafe_allow_html=True)
    else:
        st.warning("⚠️ No se detectaron bases con columnas 'estado'. Por favor, carga una base compatible.")

    # ======================================================
    # 🧠 Detección automática de base compatible (df1–df8)
    # ======================================================
    if df_auto is None:
        for i in range(1, 9):  # busca en df1 ... df8
            key = f"df{i}"
            if key in st.session_state:
                df_temp = st.session_state[key].copy()
                df_temp.columns = [c.strip() for c in df_temp.columns]  # limpia espacios
                columnas_lower = [c.lower() for c in df_temp.columns]

                # Buscar columnas que contengan la palabra "estado"
                columnas_estado = [df_temp.columns[i] for i, c in enumerate(columnas_lower) if "estado" in c]

                if len(columnas_estado) >= 2:
                    col_estado_actual = columnas_estado[0]
                    col_estado_siguiente = columnas_estado[1]
                    st.success(f"✅ Detectadas columnas de transición en **{key}**: {col_estado_actual}, {col_estado_siguiente}")
                    df_auto = df_temp.rename(columns={
                        col_estado_actual: "Estado_Actual",
                        col_estado_siguiente: "Estado_Siguiente"
                    })
                    break

                elif len(columnas_estado) == 1:
                    col_estado = columnas_estado[0]
                    st.info(f"ℹ️ Detectada una columna de estado en **{key}**: {col_estado} — se usará para calcular transiciones consecutivas.")
                    df_auto = df_temp.rename(columns={col_estado: "Estado"})
                    break

    # ======================================================
    # 📂 Si no se detectó base automáticamente
    # ======================================================
    if df_auto is None:
        file = st.file_uploader("📂 Sube una base de datos con una columna de estados (por ejemplo: 'Calidad', 'Estado de bomba', 'Condición sanitaria')", type=["csv"])
    else:
        file = None
        df = df_auto.copy()

    # ======================================================
    # 🔍 Si se cargó archivo o se detectó base automáticamente
    # ======================================================
    import pandas as pd
    import numpy as np

    df = None

    # --- Si se detectó automáticamente ---
    if df_auto is not None:
        df = df_auto.copy()

    # --- Si se subió manualmente ---
    elif "file" in locals() and file is not None:
        try:
            df = pd.read_csv(file, sep=None, engine="python")
        except Exception:
            df = pd.read_csv(file)
        df.columns = [c.strip() for c in df.columns]

    # --- Si no hay nada cargado ---
    if df is None:
        st.info("📥 Esperando que subas o se detecte una base de datos CSV para estimar el modelo de Markov.")
        st.stop()

    # ======================================================
    # 🔢 Cálculo de matriz de transición según tipo de base
    # ======================================================
    if {"Estado_Actual", "Estado_Siguiente"}.issubset(df.columns):
        estados_actual = df["Estado_Actual"].astype(str).values
        estados_siguiente = df["Estado_Siguiente"].astype(str).values
        unicos = sorted(list(set(estados_actual) | set(estados_siguiente)))
        n = len(unicos)
        P = pd.crosstab(estados_actual, estados_siguiente, normalize="index")\
            .reindex(index=unicos, columns=unicos, fill_value=0).to_numpy()

    elif "Estado" in df.columns:
        estados = df["Estado"].astype(str).values
        unicos = sorted(pd.unique(df["Estado"].astype(str)))
        n = len(unicos)
        P = np.zeros((n, n), dtype=float)
        for i in range(len(estados) - 1):
            a, b = estados[i], estados[i + 1]
            ia, ib = unicos.index(a), unicos.index(b)
            P[ia, ib] += 1
        row_sums = P.sum(axis=1, keepdims=True)
        P = np.divide(P, np.where(row_sums == 0, 1, row_sums))

    else:
        st.warning("⚠️ No se encontró ninguna columna de estado válida para construir la matriz de transición.")
        st.stop()

    # ======================================================
    # 🔢 Mostrar la matriz
    # ======================================================
    st.subheader("🔢 Matriz de transición estimada")
    st.dataframe(pd.DataFrame(P, index=unicos, columns=unicos).round(3))

    # ======================================================
    # 🔹 Simulación Monte Carlo
    # ======================================================
    pasos = st.slider("Horizonte de simulación (pasos)", 1, 50, 12)
    n_sim = st.slider("Número de simulaciones Monte Carlo", 100, 5000, 1000)
    estado_inicial = st.selectbox("Estado inicial:", unicos, index=len(unicos) - 1)

    idx_ini = unicos.index(estado_inicial)
    resultados = np.zeros((n_sim, pasos), dtype=int)

    for i in range(n_sim):
        estado = idx_ini
        for t in range(pasos):
            estado = np.random.choice(range(n), p=P[estado])
            resultados[i, t] = estado

    # ======================================================
    # 🔹 Evolución probabilística
    # ======================================================
    probs = np.zeros((pasos, n), dtype=float)
    for t in range(pasos):
        vals, counts = np.unique(resultados[:, t], return_counts=True)
        probs[t, vals] = counts / n_sim

    prob_df = pd.DataFrame(probs, columns=unicos)
    st.subheader("📊 Evolución probabilística de los estados")
    st.line_chart(prob_df)

    estado_final = unicos[np.argmax(probs[-1])]
    st.success(f"🎯 Estado más probable al final del horizonte ({pasos} pasos): **{estado_final}**")

    # ======================================================
    # 🔹 Conclusiones del modelado estocástico
    # ======================================================
    st.markdown("## 📌 Conclusiones del modelo de Markov y simulación Monte Carlo")

    def conclusion_box(text, color, icon):
        st.markdown(
            f"""
            <div style="background-color:white; border-left: 6px solid {color};
                        padding:8px; margin:6px; border-radius:5px;
                        box-shadow: 0px 2px 4px rgba(0,0,0,0.1);">
                <b>{icon}</b> {text}
            </div>
            """, unsafe_allow_html=True
        )

    conclusion_box(
        f"📘 Se estimó una matriz de transición de tamaño <b>{n}×{n}</b> con base en los estados observados.",
        "#0d6efd", ""
    )

    conclusion_box(
        f"🎲 Se ejecutaron <b>{n_sim}</b> simulaciones Monte Carlo durante <b>{pasos}</b> pasos, "
        "permitiendo visualizar la evolución probabilística del sistema en el tiempo.",
        "#6f42c1", ""
    )

    conclusion_box(
        f"🏆 El estado con mayor probabilidad al final del horizonte fue <b>{estado_final}</b>, "
        "lo que sugiere una tendencia asintótica o de equilibrio del proceso.",
        "#28a745", ""
    )

    conclusion_box(
        "🧩 Se recomienda repetir la simulación con horizontes más amplios para evaluar la estabilidad de largo plazo "
        "y utilizar esta matriz para calcular el estado estacionario del proceso de Markov.",
        "#17a2b8", ""
    )

    conclusion_box(
        "🧀 En el contexto de la gestión de la calidad de la leche, estos resultados pueden representar la probabilidad "
        "de que un tanque mantenga su clasificación de calidad o transite hacia estados de menor o mayor nivel sanitario.",
        "#ffc107", ""
    )

    # ======================================================
    # 💾 Guardar resultados en sesión para informe Word
    # ======================================================
    import io
    import matplotlib.pyplot as plt

    try:
        # --- Crear figura resumen Monte Carlo ---
        fig, ax = plt.subplots(figsize=(6, 3))
        for col in prob_df.columns:
            ax.plot(prob_df.index + 1, prob_df[col], label=col)
        ax.set_title("Evolución probabilística de los estados")
        ax.set_xlabel("Pasos de simulación")
        ax.set_ylabel("Probabilidad")
        ax.legend(title="Estados", fontsize=8, loc="best")
        ax.grid(True, alpha=0.3)

        buf = io.BytesIO()
        fig.savefig(buf, format="png", bbox_inches="tight")
        plt.close(fig)

        # Guardar en session_state
        st.session_state["img_montecarlo"] = buf.getvalue()
        st.session_state["estado_final_mc"] = estado_final
        st.session_state["n_sim_mc"] = n_sim
        st.session_state["pasos_mc"] = pasos
        st.session_state["matriz_markov"] = pd.DataFrame(P, index=unicos, columns=unicos).round(3)
        st.session_state["prob_df_mc"] = prob_df

        st.info("💾 Resultados de Monte Carlo guardados para el informe Word.")
    except Exception as e:
        st.warning(f"No se pudieron guardar los resultados de Monte Carlo: {e}")



# ===================== TAB 5: Control estadístico =====================
with tab5:
    if "best" not in st.session_state:
        st.warning("⚠️ Primero ejecute la pestaña 3 (Modelado y Pronóstico) para calcular residuales.")
    else:
        best = st.session_state.get("best_model_name", "")
        final_fit = st.session_state.get("final_fit", None)
        seasonal_periods = st.session_state.get("seasonal_periods_used", 12)


        resid = None

        try:
            if final_fit is not None:
                # Modelos que tienen fittedvalues
                resid = y - final_fit.fittedvalues.reindex(y.index).fillna(method="bfill")
            else:
                # Modelos Naive o sin objeto fit
                if "estacional" in best:
                    resid = y - y.shift(int(seasonal_periods))
                else:
                    resid = y - y.shift(1)
                resid = resid.dropna()
        except Exception as e:
            st.error(f"No se pudieron calcular residuales: {e}")


        if resid is None or resid.empty:
            st.warning("⚠️ No se pudieron calcular residuales para el modelo seleccionado.")
        else:
            # ==============================
            # 📘 Sección teórica expandible
            # ==============================
            with st.expander("📘 Fundamento teórico de los métodos de control estadístico (EWMA, Shewhart y CUSUM)"):
                st.markdown("**Fórmulas principales:**")

                # --- EWMA ---
                st.markdown("### 🔹 EWMA (Exponentially Weighted Moving Average)")
                st.latex(r"z_t = \lambda x_t + (1-\lambda)z_{t-1}, \qquad \sigma_z = \sigma \sqrt{\tfrac{\lambda}{2-\lambda}}")

                st.write("**Definiciones:**")
                st.write("- Estadístico suavizado: ", r"$z_t$ = promedio ponderado de los residuales.")
                st.write("- Residual en el instante: ", r"$x_t$.")
                st.write("- Peso del dato más reciente: ", r"$\lambda$, con $0 < \lambda \le 1$.")
                st.write("- Desviación estándar esperada: ", r"$\sigma_z$.")
                st.write("- Si", r"$z_t$", "supera los límites", r"$\pm 3\sigma_z$", "→ el proceso muestra **deriva o tendencia gradual.**")

                st.divider()

                # --- SHEWHART ---
                st.markdown("### 🔹 Shewhart (Control de medias individuales)")
                st.latex(r"UCL = \bar{x} + 3\sigma, \qquad CL = \bar{x}, \qquad LCL = \bar{x} - 3\sigma")

                st.write("**Definiciones:**")
                st.write("- Límites de control:", r"$UCL, CL, LCL$.")
                st.write("- Media histórica:", r"$\bar{x}$.")
                st.write("- Desviación estándar del proceso:", r"$\sigma$.")
                st.write("- Si algún punto excede", r"$\pm 3\sigma$", "→ **alteración puntual o dato atípico.**")

                st.divider()

                # --- CUSUM ---
                st.markdown("### 🔹 CUSUM (Cumulative Sum)")
                st.latex(r"C_t^+ = \max\{0, x_t - (\mu_0 + k) + C_{t-1}^+\}, \qquad C_t^- = \max\{0, (\mu_0 - k) - x_t + C_{t-1}^-\}")

                st.write("**Definiciones:**")
                st.write("- Sumas acumuladas positivas y negativas:", r"$C_t^+$ y $C_t^-$.")
                st.write("- Media objetivo:", r"$\mu_0$.")
                st.write("- Parámetro de referencia (sensibilidad):", r"$k$.")
                st.write("- Umbral o límite de decisión:", r"$h$.")
                st.write("- Si", r"$C_t^+ > h$", "o", r"$C_t^- > h$", "→ **cambio sostenido en la media del proceso.**")

                st.divider()

                # --- Interpretación final ---
                st.markdown("""
                ### 🧠 Interpretación aplicada
                - **EWMA** detecta desviaciones graduales (deriva del proceso).
                - **Shewhart** identifica eventos puntuales o valores atípicos.
                - **CUSUM** detecta cambios pequeños pero sostenidos en la media.
                
                💬 *En el contexto del monitoreo de calidad de leche*,  
                estos métodos permiten identificar incrementos anormales en el **conteo microbiano**, variaciones en la **temperatura del tanque** o en la **eficiencia energética**, anticipando desviaciones antes de que afecten la producción.
                """)


            # EWMA
            st.subheader("EWMA (residuales)")
            lam = st.slider("λ (EWMA)", 0.05, 0.5, 0.2, 0.05, key="lam")
            z, UCL_e, LCL_e = ewma_chart(resid, lam)
            f6, ax6 = plt.subplots(figsize=(11, 3))
            ax6.plot(z.index, z.values)
            ax6.plot(UCL_e.index, UCL_e.values, linestyle="--")
            ax6.plot(LCL_e.index, LCL_e.values, linestyle="--")
            st.pyplot(f6)
            img_ewma = fig_to_bytes(f6)

            # Shewhart
            st.subheader("Shewhart (residuales)")
            resid_s, UCL, LCL = shewhart_chart(resid)
            f5, ax5 = plt.subplots(figsize=(11, 3))
            ax5.plot(resid_s.index, resid_s.values, marker="o")
            ax5.axhline(resid_s.mean(), linestyle="--")
            ax5.axhline(UCL, color="r", linestyle="--")
            ax5.axhline(LCL, color="r", linestyle="--")
            st.pyplot(f5)
            img_shewhart = fig_to_bytes(f5)
            ooc_shewhart = resid_s[(resid_s > UCL) | (resid_s < LCL)]


            # CUSUM
            st.subheader("CUSUM (residuales)")
            sigma_res = float(resid.std(ddof=1))
            k = st.slider("k (en múltiplos de σ)", 0.1, 2.0, 0.5, 0.1, key="k") * sigma_res
            h = st.slider("h (umbral)", 1.0, 10.0, 5.0, 0.5, key="h")
            cp, cm = cusum_chart(resid, k, h)
            f7, ax7 = plt.subplots(figsize=(11, 3))
            ax7.plot(cp.index, cp.values, label="CUSUM+")
            ax7.plot(cm.index, cm.values, label="CUSUM-")
            ax7.axhline(h, linestyle="--")
            ax7.legend()
            st.pyplot(f7)
            img_cusum = fig_to_bytes(f7)

            # Guardar en session_state
            st.session_state["img_shewhart"] = img_shewhart
            st.session_state["img_ewma"] = img_ewma
            st.session_state["img_cusum"] = img_cusum

            # Conclusiones
            if len(ooc_shewhart) > 0:
                conc_control.append(f"Se detectaron {len(ooc_shewhart)} puntos fuera de control (Shewhart).")
            else:
                conc_control.append("No se detectaron puntos fuera de control en Shewhart.")
            if z.iloc[-1] > UCL_e.iloc[-1] or z.iloc[-1] < LCL_e.iloc[-1]:
                conc_control.append("EWMA indica señal de fuera de control reciente.")
            else:
                conc_control.append("EWMA no indica señal reciente de fuera de control.")

            # ----------------------------------------------------------------------
            # 🔍 Conclusiones automáticas del monitoreo (nuevo bloque)
            # ----------------------------------------------------------------------
            st.markdown("### 📊 Conclusiones automáticas del monitoreo")

            def conclusion_box(text, color):
                st.markdown(
                    f"""
                    <div style='background-color:white; border-left:6px solid {color};
                                padding:8px; margin:6px; border-radius:5px;
                                box-shadow:0px 2px 4px rgba(0,0,0,0.1);'>
                        {text}
                    </div>
                    """, unsafe_allow_html=True
                )

            import numpy as np

            # --- EWMA ---
            if np.any((z > UCL_e) | (z < LCL_e)):
                conclusion_box("🚨 El gráfico **EWMA** detecta cambios graduales sostenidos en los residuales. "
                               "El proceso podría estar mostrando deriva o nueva tendencia.", "#dc3545")
            else:
                conclusion_box("✅ El gráfico **EWMA** muestra estabilidad — sin señales de cambio estructural. "
                               "Los residuales se comportan como ruido blanco.", "#198754")

            # --- SHEWHART ---
            fuera_shewhart = np.sum((resid_s > UCL) | (resid_s < LCL))
            if fuera_shewhart > 0:
                conclusion_box(f"⚠️ El gráfico **Shewhart** presenta {fuera_shewhart} punto(s) fuera de ±3σ. "
                               "Podría existir una alteración puntual o dato atípico.", "#ffc107")
            else:
                conclusion_box("✅ El gráfico **Shewhart** indica que la variabilidad está dentro de límites normales.", "#198754")

            # --- CUSUM ---
            if np.any(cp > h) or np.any(cm > h):
                conclusion_box("🚨 El gráfico **CUSUM** muestra acumulaciones sostenidas que superan el umbral. "
                               "Esto sugiere cambios estructurales o inestabilidad en el proceso.", "#dc3545")
            else:
                conclusion_box("✅ El gráfico **CUSUM** se mantiene dentro de los límites — sin evidencia de cambio en la media.", "#198754")

            # ----------------------------------------------------------------------

            show_conclusiones("Conclusiones (Control estadístico)", conc_control)

# ===================== TAB 6: Conclusiones y Recomendaciones =====================
with tab6:
    st.subheader("📋 Conclusiones y Recomendaciones")

    conclusiones, recomendaciones = [], []

    # Consolidar todas las conclusiones previas
    conclusiones.extend(conc_exploracion + conc_stl + conc_modelo + conc_control)

    # ----------------------------------------------------------------------
    # 🔍 Conclusiones automáticas de control estadístico (vienen del TAB 5)
    # ----------------------------------------------------------------------
    st.markdown("### 📊 Conclusiones automáticas del monitoreo")
    
    # ----------------------------------------------------------------------
    # Entrada manual del usuario
    # ----------------------------------------------------------------------
    st.markdown("**Añadir notas manuales**")
    user_conc = st.text_area("Conclusiones adicionales")
    user_rec = st.text_area("Recomendaciones adicionales")
    
    if user_conc:
        conclusiones.append(user_conc)
    if user_rec:
        recomendaciones.append(user_rec)
    def conclusion_box(text, color):
        st.markdown(
            f"""
            <div style='background-color:white; border-left:6px solid {color};
                        padding:8px; margin:6px; border-radius:5px;
                        box-shadow:0px 2px 4px rgba(0,0,0,0.1);'>
                {text}
            </div>
            """, unsafe_allow_html=True
        )

    # Recuperar residuales e indicadores si existen
    resid_info = st.session_state.get("img_shewhart", None)
    if "img_shewhart" in st.session_state and "img_cusum" in st.session_state:

        # Estado de las variables clave de TAB 5
        best = st.session_state.get("best_model_name", "")
        tbl = st.session_state.get("tbl", None)

        # 1️⃣ Modelo
        if best:
            conclusion_box(f"🧩 El modelo final utilizado para control fue **{best}**.", "#0d6efd")
            conclusiones.append(f"El modelo final utilizado para control fue {best}.")

        # 2️⃣ EWMA
        if any("EWMA" in c for c in conc_control if "fuera" in c or "señal" in c):
            conclusion_box("🚨 El control **EWMA** muestra posible deriva o tendencia. Se recomienda revisar calibración del proceso.",
                           "#dc3545")
            recomendaciones.append("Revisar la calibración o tendencia gradual detectada en el control EWMA.")
        else:
            conclusion_box("✅ El control **EWMA** no presenta señales de deriva — proceso estable.",
                           "#198754")

        # 3️⃣ Shewhart
        if any("fuera de control" in c for c in conc_control):
            conclusion_box("⚠️ El gráfico **Shewhart** indica puntos fuera de ±3σ. Puede existir una alteración puntual.",
                           "#ffc107")
            recomendaciones.append("Investigar causas específicas de los puntos fuera de control (Shewhart).")
        else:
            conclusion_box("✅ El gráfico **Shewhart** mantiene los puntos dentro de ±3σ — sin alteraciones relevantes.",
                           "#198754")

        # 4️⃣ CUSUM
        if any("CUSUM" in c or "acumulaciones" in c for c in conc_control):
            conclusion_box("🚨 El gráfico **CUSUM** muestra acumulaciones sostenidas — posible cambio estructural.",
                           "#dc3545")
            recomendaciones.append("Analizar causas de cambios estructurales o variaciones persistentes en el proceso.")
        else:
            conclusion_box("✅ El gráfico **CUSUM** no evidencia cambios acumulativos significativos.",
                           "#198754")

    else:
        st.info("ℹ️ No se han generado resultados de control estadístico aún.")

    # ----------------------------------------------------------------------
    # Recomendaciones automáticas (según precisión y pronóstico)
    # ----------------------------------------------------------------------
    if "tbl" in st.session_state and "best" in st.session_state:
        m = st.session_state["tbl"].loc[st.session_state["best"]]["MAPE%"]
        if pd.notna(m) and m < 10:
            recomendaciones.append("Implementar el modelo seleccionado en decisiones operativas.")
        elif pd.notna(m) and m < 20:
            recomendaciones.append("Usar el modelo con validación de experto y monitoreo cercano.")
        else:
            recomendaciones.append("Revisar calidad de datos y probar otros modelos.")

    if "future" in st.session_state:
        fmean, ymean = float(st.session_state["future"].mean()), float(y.mean())
        if fmean > ymean:
            recomendaciones.append("Planificar recursos para atender mayor demanda futura.")
        elif fmean < ymean:
            recomendaciones.append("Ajustar planificación para evitar sobreproducción.")
        else:
            recomendaciones.append("Mantener el nivel operativo actual y continuar monitoreando la tendencia.")

    # Guardar resultados en session_state
    st.session_state["conclusiones"] = conclusiones
    st.session_state["recomendaciones"] = recomendaciones

    # ----------------------------------------------------------------------
    # Visualización final consolidada
    # ----------------------------------------------------------------------
    show_conclusiones("Conclusiones (Consolidadas)", conclusiones)
    show_recomendaciones("Recomendaciones (Consolidadas)", recomendaciones)

# ===================== TAB 7: Exportar Word =====================
with tab7:
    st.subheader("💾 Exportar informe en Word")

    if "tbl" not in st.session_state or "best" not in st.session_state:
        st.warning("Primero ejecute la pestaña 3 para generar resultados.")
    else:
        from docx import Document
        from docx.shared import Inches
        from datetime import datetime
        import tempfile, io

        def safe_add_pic(doc, img_bytes, title):
            """Agrega imagen al Word si existe"""
            if img_bytes:
                doc.add_paragraph(title, style="List Bullet")
                doc.add_picture(io.BytesIO(img_bytes), width=Inches(5))
                doc.add_paragraph("")

        # Crear documento
        doc = Document()
        doc.add_heading("Informe – Super Laboratorio (Unidad 4)", 0)
        doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph(f"Modelo seleccionado: {st.session_state['best']}")
        doc.add_paragraph("")

        # ----------------------------------------------------------
        # 📊 DATOS BÁSICOS
        # ----------------------------------------------------------
        doc.add_heading("Datos", level=1)
        doc.add_paragraph(f"Número de observaciones: {len(y)}")
        doc.add_paragraph(f"Frecuencia temporal: {y.index.freqstr or 'no inferida'}")
        if "freq_detected" in st.session_state:
            doc.add_paragraph(f"Frecuencia detectada: {st.session_state['freq_detected']}")
        doc.add_paragraph("")

        # ----------------------------------------------------------
        # 📈 DESEMPEÑO DEL MODELO
        # ----------------------------------------------------------
        doc.add_heading("Desempeño del modelo", level=1)
        best = st.session_state["best"]
        tbl = st.session_state["tbl"].loc[best]
        for k, v in tbl.items():
            doc.add_paragraph(f"{k}: {v}")
        doc.add_paragraph("")

        # ----------------------------------------------------------
        # 📉 GRÁFICOS
        # ----------------------------------------------------------
        doc.add_heading("Gráficas", level=1)
        safe_add_pic(doc, img_series, "Serie temporal observada")
        safe_add_pic(doc, st.session_state.get("img_fit"), "Ajuste y conjunto de prueba")
        safe_add_pic(doc, st.session_state.get("img_forecast"), "Pronóstico futuro")
        safe_add_pic(doc, st.session_state.get("img_stl"), "Descomposición estacional (STL)")
        safe_add_pic(doc, st.session_state.get("img_shewhart"), "Carta de control Shewhart (residuales)")
        safe_add_pic(doc, st.session_state.get("img_ewma"), "Carta de control EWMA")
        safe_add_pic(doc, st.session_state.get("img_cusum"), "Carta de control CUSUM")
        doc.add_paragraph("")

        # ----------------------------------------------------------
        # 🎲 SIMULACIÓN MONTE CARLO
        # ----------------------------------------------------------
        doc.add_heading("Simulación Monte Carlo", level=1)
        if "img_montecarlo" in st.session_state:
            safe_add_pic(doc, st.session_state["img_montecarlo"], "Evolución probabilística de los estados")
            doc.add_paragraph(f"Se realizaron {st.session_state['n_sim_mc']} simulaciones durante {st.session_state['pasos_mc']} pasos.")
            doc.add_paragraph(f"El estado más probable al final del horizonte fue: {st.session_state['estado_final_mc']}.")
            doc.add_paragraph("")
            doc.add_paragraph("Matriz de transición utilizada:")
            tabla_mc = st.session_state["matriz_markov"]
            for fila in tabla_mc.to_string().split("\n"):
                doc.add_paragraph(fila, style="No Spacing")
        else:
            doc.add_paragraph("No se generó ninguna simulación Monte Carlo en esta sesión.")

        # ----------------------------------------------------------
        # 🧠 CONCLUSIONES AUTOMÁTICAS
        # ----------------------------------------------------------
        doc.add_heading("Conclusiones", level=1)
        if "conclusiones" in st.session_state:
            for c in st.session_state.get("conclusiones", []):
                doc.add_paragraph(c, style="List Bullet")
        else:
            doc.add_paragraph("No se registraron conclusiones automáticas.")

        # ----------------------------------------------------------
        # 🧩 RECOMENDACIONES
        # ----------------------------------------------------------
        doc.add_heading("Recomendaciones", level=1)
        if "recomendaciones" in st.session_state:
            for r in st.session_state.get("recomendaciones", []):
                doc.add_paragraph(r, style="List Bullet")
        else:
            doc.add_paragraph("No se registraron recomendaciones.")

        # ----------------------------------------------------------
        # 🧾 CONTROL ESTADÍSTICO (Resumen)
        # ----------------------------------------------------------
        doc.add_heading("Control Estadístico", level=1)
        if "img_shewhart" in st.session_state or "img_ewma" in st.session_state or "img_cusum" in st.session_state:
            doc.add_paragraph("El sistema de control estadístico permitió evaluar la estabilidad de los residuales mediante:")
            doc.add_paragraph("• Carta de Shewhart: detección de puntos fuera de ±3σ.")
            doc.add_paragraph("• Carta EWMA: detección de tendencias o derivas graduales.")
            doc.add_paragraph("• Carta CUSUM: detección de acumulaciones sostenidas.")
        else:
            doc.add_paragraph("No se generaron gráficos de control estadístico en esta sesión.")

        # ----------------------------------------------------------
        # 💾 EXPORTACIÓN FINAL
        # ----------------------------------------------------------
        tmpf = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmpf.name)

        with open(tmpf.name, "rb") as f:
            st.download_button(
                "⬇️ Descargar informe (.docx)",
                f,
                file_name=f"informe_unidad4_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            )













